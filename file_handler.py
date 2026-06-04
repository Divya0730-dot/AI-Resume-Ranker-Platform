import os
import uuid
from datetime import datetime
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
import PyPDF2
from flask import Response, make_response
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import io
from config import Config

class FileHandler:
    def __init__(self):
        self.upload_folder = Config.UPLOAD_FOLDER
        self.allowed_extensions = Config.ALLOWED_EXTENSIONS
        
        # Create upload directory if it doesn't exist
        os.makedirs(self.upload_folder, exist_ok=True)
    
    def allowed_file(self, filename):
        """Check if file extension is allowed"""
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in self.allowed_extensions
    
    def save_file(self, file, username):
        """Save uploaded file with secure filename"""
        if file and self.allowed_file(file.filename):
            # Create secure filename with timestamp
            original_filename = secure_filename(file.filename)
            name, ext = os.path.splitext(original_filename)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_id = str(uuid.uuid4())[:8]
            
            filename = f"{username}_{name}_{timestamp}_{unique_id}{ext}"
            file_path = os.path.join(self.upload_folder, filename)
            
            file.save(file_path)
            return file_path
        
        raise ValueError("Invalid file")
    
    def extract_text_from_pdf(self, file_path):
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error extracting DOCX text: {e}")
            return ""
    
    def extract_text(self, file_path):
        """Extract text from file based on extension"""
        if not os.path.exists(file_path):
            raise FileNotFoundError("File not found")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            # Handle text files for testing
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read().strip()
            except Exception as e:
                print(f"Error reading text file: {e}")
                return ""
        else:
            # Try to read as text file if extension is unknown but file exists
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read().strip()
                    if content:  # If we successfully read content, return it
                        return content
            except:
                pass
            raise ValueError("Unsupported file format")
    
    def delete_file(self, file_path):
        """Delete file from filesystem"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                return True
        except Exception as e:
            print(f"Error deleting file: {e}")
        return False
    
    def convert_to_pdf(self, text_content, original_filename):
        """Convert text content to PDF and return as response"""
        try:
            # Create PDF in memory
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, 
                                  rightMargin=72, leftMargin=72, 
                                  topMargin=72, bottomMargin=18)
            
            # Define styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                leading=14
            )
            
            # Build PDF content
            story = []
            
            # Add title
            filename_without_ext = os.path.splitext(original_filename)[0]
            story.append(Paragraph(f"Resume: {filename_without_ext}", title_style))
            story.append(Spacer(1, 12))
            
            # Add content paragraphs
            paragraphs = text_content.split('\n')
            for para in paragraphs:
                if para.strip():
                    # Escape HTML characters and handle special characters
                    para_escaped = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(para_escaped, normal_style))
                else:
                    story.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(story)
            
            # Get PDF data
            pdf_data = buffer.getvalue()
            buffer.close()
            
            # Create response
            response = make_response(pdf_data)
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = f'attachment; filename="{filename_without_ext}.pdf"'
            
            return response
            
        except Exception as e:
            print(f"Error converting to PDF: {e}")
            return make_response({"error": "Failed to convert to PDF"}, 500)
    
    def convert_to_docx(self, text_content, original_filename):
        """Convert text content to DOCX and return as response"""
        try:
            # Create DOCX in memory
            doc = Document()
            
            # Add title
            filename_without_ext = os.path.splitext(original_filename)[0]
            title = doc.add_heading(f'Resume: {filename_without_ext}', 0)
            title.alignment = 1  # Center alignment
            
            # Add content
            paragraphs = text_content.split('\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
                else:
                    doc.add_paragraph('')  # Empty paragraph for spacing
            
            # Save to memory buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Create response
            response = make_response(buffer.getvalue())
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = f'attachment; filename="{filename_without_ext}.docx"'
            
            return response
            
        except Exception as e:
            print(f"Error converting to DOCX: {e}")
            return make_response({"error": "Failed to convert to DOCX"}, 500)