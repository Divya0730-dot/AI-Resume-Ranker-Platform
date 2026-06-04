from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify, send_file
from flask_cors import CORS
import bcrypt
import re
import os
from datetime import datetime, timezone
from config import Config
from services.database import Database
from services.ai_engine import AIEngine
from services.file_handler import FileHandler
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Pt
import PyPDF2

app = Flask(__name__)
app.config.from_object(Config)
CORS(app)

# Initialize services
db = Database()
ai_engine = AIEngine()
file_handler = FileHandler()

# Password validation regex
PASSWORD_REGEX = r'^(?=.*[a-z])(?=.*[A-Z])(?=.*\d)(?=.*[@$!%*?&])[A-Za-z\d@$!%*?&]{8,}$'

# ==================== FRONTEND ROUTES ====================
@app.route('/')
def root():
    return render_template('landing.html')

@app.route('/landing')
def landing():
    return render_template('landing.html')

@app.route('/home')
def home():
    return render_template('landing.html')

@app.route('/register')
def register():
    return render_template('register.html')

@app.route('/login')
def login():
    return render_template('login.html')

@app.route('/login-debug')
def login_debug():
    return render_template('login_debug.html')

@app.route('/simple-test')
def simple_test():
    return render_template('simple_login_test.html')

@app.route('/admin-login')
def admin_login():
    return render_template('admin_login.html')

@app.route('/dashboard')
def dashboard():
    if 'user' not in session:
        return redirect(url_for('login'))
    
    # Check user role
    user = db.get_user(session['user'])
    if not user or user.get('role') != 'Recruiter':
        flash("Access denied. Recruiter role required.")
        return redirect(url_for('login'))
    
    return render_template('dashboard.html', username=session['user'])

@app.route('/admin-dashboard')
def admin_dashboard():
    if 'admin' not in session:
        flash("Access denied. Admin login required.")
        return redirect(url_for('admin_login'))
    
    return render_template('admin_dashboard.html', admin_username=session['admin'])

# ==================== API ROUTES ====================

@app.route('/api/register', methods=['POST'])
def api_register():
    try:
        data = request.get_json()
        
        # Check if all fields are provided
        required_fields = ['fullname', 'username', 'password', 'confirm_password', 'role']
        for field in required_fields:
            if not data.get(field) or not data.get(field).strip():
                return jsonify({"error": "All fields are required"}), 400
        
        fullname = data['fullname'].strip()
        username = data['username'].strip()
        password = data['password']
        confirm_password = data['confirm_password']
        role = data['role'].strip()
        
        # Role validation
        valid_roles = ['Recruiter']
        if role not in valid_roles:
            return jsonify({"error": "Invalid role selected"}), 400
        
        # Username uniqueness check
        if db.user_exists(username):
            return jsonify({"error": "User already exists"}), 400
        
        # Password strength validation
        if not re.match(PASSWORD_REGEX, password):
            return jsonify({"error": "Password must contain uppercase, lowercase, number, special character and be at least 8 characters long"}), 400
        
        # Confirm password validation
        if password != confirm_password:
            return jsonify({"error": "Password didn't match"}), 400
        
        # Hash password with consistent salt rounds
        hashed_password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt(rounds=12))
        
        # Save user to database
        user_data = {
            "full_name": fullname,
            "username": username,
            "password": hashed_password,
            "role": role,
            "created_at": datetime.now(timezone.utc)
        }
        
        db.create_user(user_data)
        return jsonify({"message": "User registered successfully"}), 201
        
    except Exception as e:
        print(f"Registration error: {e}")
        return jsonify({"error": "Internal server error"}), 500

@app.route('/api/login', methods=['POST'])
def api_login():
    """Recruiter login API - ONLY for recruiters"""
    try:
        # Handle both JSON and form data
        if request.is_json:
            data = request.get_json()
        else:
            data = {
                'username': request.form.get('username'),
                'password': request.form.get('password')
            }
            
        if not data:
            return jsonify({"error": "Invalid username or password"}), 401
            
        username = data.get('username') or ''
        password = data.get('password') or ''
        
        # Safely strip username if it's a string
        if isinstance(username, str):
            username = username.strip()
        else:
            username = ''
            
        if not username or not password:
            return jsonify({"error": "Invalid username or password"}), 401
        
        # RECRUITER LOGIN ONLY - No admin login allowed here
        # Check for recruiter in database
        user = db.get_user(username)
        
        if not user:
            return jsonify({"error": "Invalid username or password"}), 401
        
        # STRICT: Only allow Recruiter role - reject all others including admin
        if user.get('role') != 'Recruiter':
            return jsonify({"error": "Invalid username or password"}), 401
        
        # Compare entered password with hashed password using bcrypt.checkpw()
        stored_password = user['password']
        
        # Handle both bytes and string stored passwords for compatibility
        if isinstance(stored_password, str):
            stored_password = stored_password.encode('utf-8')
        
        # Use bcrypt.checkpw() for proper password comparison
        try:
            password_match = bcrypt.checkpw(password.encode('utf-8'), stored_password)
            
            if password_match:
                session['user'] = username
                session['role'] = 'recruiter'
                return jsonify({"message": "Login successful", "role": "recruiter"}), 200
            else:
                return jsonify({"error": "Invalid username or password"}), 401
                
        except Exception as bcrypt_error:
            return jsonify({"error": "Invalid username or password"}), 401
            
    except Exception as e:
        print(f"Login error: {e}")
        return jsonify({"error": "Invalid username or password"}), 401

@app.route('/api/admin-login', methods=['POST'])
def api_admin_login():
    """Admin login API - ONLY for admins"""
    try:
        data = request.get_json()
        username = data.get('username', '').strip()
        password = data.get('password', '')
        
        if not username or not password:
            return jsonify({"error": "Invalid credentials"}), 401
        
        # Method 1: Check predefined admin credentials
        if username == Config.ADMIN_USERNAME and password == Config.ADMIN_PASSWORD:
            session['admin'] = username
            session['role'] = 'admin'
            return jsonify({"message": "Admin login successful"}), 200
        
        # Method 2: Check database for admin users
        user = db.get_user(username)
        
        if user and user.get('role') == 'admin':
            # Verify password for database admin users
            stored_password = user['password']
            
            # Handle both bytes and string stored passwords for compatibility
            if isinstance(stored_password, str):
                stored_password = stored_password.encode('utf-8')
            
            try:
                password_match = bcrypt.checkpw(password.encode('utf-8'), stored_password)
                
                if password_match:
                    session['admin'] = username
                    session['role'] = 'admin'
                    return jsonify({"message": "Admin login successful"}), 200
                else:
                    return jsonify({"error": "Invalid credentials"}), 401
                    
            except Exception as bcrypt_error:
                return jsonify({"error": "Invalid credentials"}), 401
        
        # If neither predefined admin nor database admin, reject
        return jsonify({"error": "Invalid credentials"}), 401
            
    except Exception as e:
        print(f"Admin login error: {e}")
        return jsonify({"error": "Invalid credentials"}), 401

@app.route('/api/clear-session', methods=['POST'])
def api_clear_session():
    """Clear current ranking session to start fresh"""
    try:
        if 'user' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        # Clear the current ranking session
        if 'current_ranking_session' in session:
            session.pop('current_ranking_session')
        
        # Mark old resumes as not current session
        db.clear_old_session_resumes(session['user'])
        
        return jsonify({"message": "Session cleared successfully"}), 200
        
    except Exception as e:
        print(f"Clear session error: {e}")
        return jsonify({"error": "Failed to clear session"}), 500

@app.route('/api/upload-resumes', methods=['POST'])
def api_upload_resumes():
    try:
        if 'user' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        if 'resumes' not in request.files:
            return jsonify({"error": "No files uploaded"}), 400
        
        files = request.files.getlist('resumes')
        if not files or all(file.filename == '' for file in files):
            return jsonify({"error": "No files selected"}), 400
        
        # Generate a unique session ID for this ranking batch
        import uuid
        ranking_session_id = str(uuid.uuid4())
        
        # Store session ID in user session for tracking
        session['current_ranking_session'] = ranking_session_id
        
        uploaded_files = []
        
        for file in files:
            if file.filename == '':
                continue
                
            # Validate file format
            if not file_handler.allowed_file(file.filename):
                return jsonify({"error": f"Invalid file format for {file.filename}. Only PDF and DOCX allowed"}), 400
            
            # Save file and extract text
            try:
                file_path = file_handler.save_file(file, session['user'])
                extracted_text = file_handler.extract_text(file_path)
            except Exception as e:
                print(f"File processing error for {file.filename}: {e}")
                return jsonify({"error": f"Failed to process file {file.filename}: {str(e)}"}), 500
            
            # Extract candidate name from resume text (improved method)
            candidate_name = ai_engine.extract_candidate_name(file_path, session['user'], extracted_text)
            
            # Save resume data to database with session tracking
            resume_data = {
                "username": session['user'],
                "filename": file.filename,
                "resume_path": file_path,
                "extracted_text": extracted_text,
                "candidate_name": candidate_name,  # Add candidate name
                "uploaded_at": datetime.now(timezone.utc),
                "ranking_session_id": ranking_session_id,  # Track session
                "is_current_session": True  # Mark as current session
            }
            
            try:
                db.save_resume(resume_data)
                uploaded_files.append({
                    "filename": file.filename,
                    "path": file_path,
                    "session_id": ranking_session_id
                })
            except Exception as e:
                print(f"Database save error for {file.filename}: {e}")
                return jsonify({"error": f"Failed to save {file.filename} to database"}), 500
        
        return jsonify({
            "message": f"Successfully uploaded {len(uploaded_files)} resume(s)",
            "uploaded_files": uploaded_files,
            "session_id": ranking_session_id
        }), 201
        
    except Exception as e:
        print(f"Upload error: {e}")
        return jsonify({"error": "Failed to upload resumes"}), 500
    except Exception as e:
        print(f"Upload error: {e}")
        return jsonify({"error": "Failed to upload resumes"}), 500

@app.route('/api/generate-job-description', methods=['POST'])
def api_generate_job_description():
    try:
        if 'user' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        data = request.get_json()
        job_position = data.get('job_position', '').strip()
        job_title = data.get('job_title', '').strip()
        
        if not job_position or not job_title:
            return jsonify({"error": "Job position and title are required"}), 400
        
        # Generate job description using AI engine
        job_description = ai_engine.generate_job_description(job_position, job_title)
        
        return jsonify({"job_description": job_description}), 200
        
    except Exception as e:
        print(f"Job description generation error: {e}")
        return jsonify({"error": "Failed to generate job description"}), 500

@app.route('/api/view-resume/<filename>', methods=['GET'])
def api_view_resume(filename):
    try:
        if 'user' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        # Get resume content from database by filename
        resume = db.get_resume_by_filename(session['user'], filename)
        if not resume:
            return jsonify({"error": "Resume not found"}), 404
        
        return jsonify({"content": resume.get('extracted_text', '')}), 200
        
    except Exception as e:
        print(f"Resume view error: {e}")
        return jsonify({"error": "Failed to load resume"}), 500

@app.route('/api/download-resume/<filename>/<format>', methods=['GET'])
def api_download_resume(filename, format):
    try:
        if 'user' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        # Validate format first
        if format.lower() not in ['pdf', 'docx']:
            return jsonify({"error": "Invalid format. Use 'pdf' or 'docx'"}), 400
        
        # Get resume from database
        resume = db.get_resume_by_filename(session['user'], filename)
        if not resume:
            return jsonify({"error": "Resume not found"}), 404
        
        resume_path = resume.get('resume_path', '')
        extracted_text = resume.get('extracted_text', '')
        
        if format.lower() == 'pdf':
            return file_handler.convert_to_pdf(extracted_text, filename)
        elif format.lower() == 'docx':
            return file_handler.convert_to_docx(extracted_text, filename)
        
    except Exception as e:
        print(f"Resume download error: {e}")
        return jsonify({"error": "Failed to download resume"}), 500

@app.route('/api/print-resume/<filename>', methods=['GET'])
def api_print_resume(filename):
    try:
        if 'user' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        # Get resume content from database by filename
        resume = db.get_resume_by_filename(session['user'], filename)
        if not resume:
            return jsonify({"error": "Resume not found"}), 404
        
        # Return formatted content for printing
        return jsonify({
            "content": resume.get('extracted_text', ''),
            "filename": filename,
            "candidate_name": resume.get('candidate_name', 'Unknown Candidate')
        }), 200
        
    except Exception as e:
        print(f"Resume print error: {e}")
        return jsonify({"error": "Failed to prepare resume for printing"}), 500

@app.route('/api/rank-resumes', methods=['POST'])
def api_rank_resumes():
    try:
        if 'user' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        data = request.get_json()
        job_title = data.get('job_title', '').strip()
        job_description = data.get('job_description', '').strip()
        job_position = data.get('job_position', '').strip()
        
        if not job_description:
            return jsonify({"error": "Job description is required"}), 400
        
        if not job_title:
            return jsonify({"error": "Job title is required"}), 400
        
        # Check if there's a current ranking session
        current_session_id = session.get('current_ranking_session')
        resumes = []
        mode = "session"
        
        if current_session_id:
            # Get resumes from current session (session isolation)
            resumes = db.get_session_resumes(session['user'], current_session_id)
            print(f"Found {len(resumes)} resumes from current session {current_session_id}")
            
            if resumes:
                mode = "session"
            else:
                print(f"⚠️ No resumes found for current session {current_session_id}")
                # Session exists but no resumes - user needs to upload
                return jsonify({"error": "No resumes found in current session. Please upload resumes first"}), 400
        else:
            print("⚠️ No current ranking session found")
            # No session - user needs to upload resumes to create a session
            return jsonify({"error": "No active session found. Please upload resumes first to start a new ranking session"}), 400
        
        if not resumes:
            return jsonify({"error": "No resumes found in current session. Please upload resumes first"}), 400
        
        print(f"Ranking {len(resumes)} resumes for job: {job_position}")
        
        # Rank resumes using AI engine
        ranked_results = ai_engine.rank_resumes(resumes, job_description, job_title)
        
        # Transform the results to match frontend expectations
        transformed_results = []
        for i, result in enumerate(ranked_results):
            # Get the original resume data to extract filename
            original_resume = resumes[i] if i < len(resumes) else {}
            
            # Extract filename from the original resume data
            filename = original_resume.get('filename', 'Unknown Resume')
            if not filename and 'resume_path' in original_resume:
                filename = os.path.basename(original_resume['resume_path'])
            
            # Map skill_recommendations to recommended_skills and extract missing skills
            skill_recommendations = result.get('skill_recommendations', [])
            existing_skills = result.get('existing_skills', [])
            
            # skill_recommendations from AI engine are the missing/recommended skills
            recommended_skills = skill_recommendations
            missing_skills = skill_recommendations  # Same as recommended since these are what's missing
            
            # Get enhanced skill analysis data directly from result
            hiring_strategies = result.get('hiring_strategies', [])
            development_plan = result.get('development_plan', {})
            skill_gap_analysis = result.get('skill_gap_analysis', {})
            
            transformed_result = {
                'filename': filename,
                'match_score': result.get('score', 0),
                'recommended_skills': recommended_skills,
                'missing_skills': missing_skills,  # Add missing skills separately
                'existing_skills': existing_skills,
                'rank': result.get('rank', i + 1),
                'candidate_name': result.get('candidate_name', 'Unknown'),
                'skills_summary': result.get('skills_summary', {}),
                'resume_path': original_resume.get('resume_path', ''),
                'extracted_text': original_resume.get('extracted_text', ''),
                'session_id': current_session_id or 'legacy',
                'job_position': job_position,
                'job_title': job_title,
                # Enhanced skill analysis data
                'hiring_strategies': hiring_strategies,
                'development_plan': development_plan,
                'skill_gap_analysis': skill_gap_analysis
            }
            transformed_results.append(transformed_result)
        
        # Save ranking results to database for admin tracking
        ranking_data = {
            "username": session['user'],
            "session_id": current_session_id or 'legacy',
            "job_position": job_position,
            "job_title": job_title,
            "job_description": job_description,
            "ranked_at": datetime.now(timezone.utc),
            "results_count": len(transformed_results),
            "mode": mode
        }
        db.save_ranking_session(ranking_data)
        
        # Update each resume with its score for admin dashboard
        for i, result in enumerate(transformed_results):
            original_resume = resumes[i] if i < len(resumes) else {}
            resume_id = str(original_resume.get('_id', ''))
            if resume_id:
                db.update_resume_score(resume_id, result['match_score'], result.get('existing_skills', []))
        
        return jsonify({
            "rankings": transformed_results,
            "session_id": current_session_id or 'legacy',
            "mode": mode,
            "job_info": {
                "position": job_position,
                "title": job_title,
                "description": job_description
            }
        }), 200
        
    except Exception as e:
        print(f"Ranking error: {e}")
        return jsonify({"error": "Failed to rank resumes"}), 500

@app.route('/api/logout', methods=['POST'])
def api_logout():
    session.pop('user', None)
    session.pop('admin', None)
    session.pop('role', None)
    return jsonify({"message": "Logged out successfully"}), 200

# Admin API endpoints
@app.route('/api/admin/users', methods=['GET'])
def api_admin_get_users():
    try:
        if 'admin' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        users = db.get_all_users()
        # Remove password field and convert ObjectId to string for security and JSON serialization
        for user in users:
            user.pop('password', None)
            if '_id' in user:
                user['_id'] = str(user['_id'])
            
        return jsonify({"users": users}), 200
        
    except Exception as e:
        print(f"Admin get users error: {e}")
        return jsonify({"error": "Failed to fetch users"}), 500

@app.route('/api/admin/resumes', methods=['GET'])
def api_admin_get_resumes():
    try:
        if 'admin' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        resumes = db.get_all_resumes()
        # Convert ObjectId to string for JSON serialization
        for resume in resumes:
            if '_id' in resume:
                resume['_id'] = str(resume['_id'])
                
        return jsonify({"resumes": resumes}), 200
        
    except Exception as e:
        print(f"Admin get resumes error: {e}")
        return jsonify({"error": "Failed to fetch resumes"}), 500

@app.route('/api/admin/resumes/clear-all', methods=['DELETE'])
def api_admin_clear_all_resumes():
    """Admin route to delete all resumes"""
    try:
        if 'admin' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        # Get all resumes
        resumes = db.get_all_resumes()
        
        deleted_count = 0
        for resume in resumes:
            resume_path = resume.get('resume_path', '')
            
            # Delete file from filesystem if it exists
            if resume_path and os.path.exists(resume_path):
                try:
                    os.remove(resume_path)
                    print(f"Deleted file: {resume_path}")
                except Exception as e:
                    print(f"Error deleting file: {e}")
            
            # Delete from database
            resume_id = str(resume.get('_id', ''))
            if resume_id:
                result = db.delete_resume(resume_id)
                if result:
                    deleted_count += 1
        
        return jsonify({"message": f"Successfully deleted {deleted_count} resume(s)"}), 200
        
    except Exception as e:
        print(f"Admin clear all resumes error: {e}")
        return jsonify({"error": "Failed to clear resumes"}), 500

@app.route('/api/admin/rankings', methods=['GET'])
def api_admin_get_rankings():
    try:
        if 'admin' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        rankings = db.get_all_rankings()
        # Convert ObjectId to string for JSON serialization
        for ranking in rankings:
            if '_id' in ranking:
                ranking['_id'] = str(ranking['_id'])
                
        return jsonify({"rankings": rankings}), 200
        
    except Exception as e:
        print(f"Admin get rankings error: {e}")
        return jsonify({"error": "Failed to fetch rankings"}), 500

@app.route('/api/admin/user/<username>/toggle', methods=['POST'])
def api_admin_toggle_user(username):
    try:
        if 'admin' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        data = request.get_json()
        active = data.get('active', True)
        
        result = db.update_user_status(username, active)
        if result:
            status = "activated" if active else "deactivated"
            return jsonify({"message": f"User {username} {status} successfully"}), 200
        else:
            return jsonify({"error": "User not found"}), 404
        
    except Exception as e:
        print(f"Admin toggle user error: {e}")
        return jsonify({"error": "Failed to update user status"}), 500

@app.route('/api/admin/stats', methods=['GET'])
def api_admin_get_stats():
    try:
        if 'admin' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        stats = db.get_system_stats()
        return jsonify({"stats": stats}), 200
        
    except Exception as e:
        print(f"Admin get stats error: {e}")
        return jsonify({"error": "Failed to fetch statistics"}), 500

@app.route('/api/admin/resume/<resume_id>/download', methods=['GET'])
def api_admin_download_resume(resume_id):
    """Admin route to download resume by ID"""
    try:
        if 'admin' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        # Get resume from database by ID
        resume = db.get_resume_by_id(resume_id)
        if not resume:
            return jsonify({"error": "Resume not found"}), 404
        
        resume_path = resume.get('resume_path', '')
        filename = resume.get('filename', 'resume.pdf')
        
        # Check if file exists
        if not os.path.exists(resume_path):
            return jsonify({"error": "Resume file not found on server"}), 404
        
        # Send the actual file
        return send_file(resume_path, as_attachment=True, download_name=filename)
        
    except Exception as e:
        print(f"Admin download resume error: {e}")
        return jsonify({"error": "Failed to download resume"}), 500

@app.route('/api/admin/resume/<resume_id>/delete', methods=['DELETE'])
def api_admin_delete_resume(resume_id):
    """Admin route to delete resume by ID"""
    try:
        if 'admin' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        # Get resume from database by ID
        resume = db.get_resume_by_id(resume_id)
        if not resume:
            return jsonify({"error": "Resume not found"}), 404
        
        resume_path = resume.get('resume_path', '')
        
        # Delete file from filesystem if it exists
        if resume_path and os.path.exists(resume_path):
            try:
                os.remove(resume_path)
                print(f"Deleted file: {resume_path}")
            except Exception as e:
                print(f"Error deleting file: {e}")
        
        # Delete from database
        result = db.delete_resume(resume_id)
        
        if result:
            return jsonify({"message": "Resume deleted successfully"}), 200
        else:
            return jsonify({"error": "Failed to delete resume from database"}), 500
        
    except Exception as e:
        print(f"Admin delete resume error: {e}")
        return jsonify({"error": "Failed to delete resume"}), 500

@app.route('/api/admin/resume/<resume_id>/save-pdf', methods=['GET'])
def api_admin_save_resume_as_pdf(resume_id):
    """Admin route to save resume as PDF"""
    try:
        if 'admin' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        # Get resume from database by ID
        resume = db.get_resume_by_id(resume_id)
        if not resume:
            return jsonify({"error": "Resume not found"}), 404
        
        resume_path = resume.get('resume_path', '')
        filename = resume.get('filename', 'resume')
        candidate_name = resume.get('candidate_name', 'Unknown')
        extracted_text = resume.get('extracted_text', 'No content available')
        
        # If file exists and is PDF, send it directly
        if resume_path and os.path.exists(resume_path) and resume_path.lower().endswith('.pdf'):
            base_name = os.path.splitext(filename)[0]
            return send_file(resume_path, as_attachment=True, download_name=f"{base_name}.pdf")
        
        # Otherwise, create PDF from extracted text
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Add title
            title = Paragraph(f"<b>{candidate_name}</b>", styles['Title'])
            story.append(title)
            story.append(Spacer(1, 12))
            
            # Add content
            for line in extracted_text.split('\n'):
                if line.strip():
                    # Escape special characters for reportlab
                    safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    p = Paragraph(safe_line, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 6))
            
            doc.build(story)
            buffer.seek(0)
            
            base_name = os.path.splitext(filename)[0]
            return send_file(buffer, as_attachment=True, download_name=f"{base_name}.pdf", mimetype='application/pdf')
        except Exception as e:
            print(f"PDF conversion error: {e}")
            return jsonify({"error": f"Failed to convert to PDF: {str(e)}"}), 500
        
    except Exception as e:
        print(f"Admin save as PDF error: {e}")
        return jsonify({"error": f"Failed to save as PDF: {str(e)}"}), 500

@app.route('/api/admin/resume/<resume_id>/save-doc', methods=['GET'])
def api_admin_save_resume_as_doc(resume_id):
    """Admin route to save resume as DOC"""
    try:
        if 'admin' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        # Get resume from database by ID
        resume = db.get_resume_by_id(resume_id)
        if not resume:
            return jsonify({"error": "Resume not found"}), 404
        
        resume_path = resume.get('resume_path', '')
        filename = resume.get('filename', 'resume')
        candidate_name = resume.get('candidate_name', 'Unknown')
        extracted_text = resume.get('extracted_text', 'No content available')
        
        # If file exists and is DOCX, send it directly
        if resume_path and os.path.exists(resume_path) and resume_path.lower().endswith('.docx'):
            base_name = os.path.splitext(filename)[0]
            return send_file(resume_path, as_attachment=True, download_name=f"{base_name}.docx")
        
        # Otherwise, create DOCX from extracted text
        try:
            buffer = BytesIO()
            document = Document()
            
            # Add title
            document.add_heading(candidate_name, 0)
            
            # Add content
            for line in extracted_text.split('\n'):
                if line.strip():
                    document.add_paragraph(line)
            
            document.save(buffer)
            buffer.seek(0)
            
            base_name = os.path.splitext(filename)[0]
            return send_file(buffer, as_attachment=True, download_name=f"{base_name}.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        except Exception as e:
            print(f"DOCX conversion error: {e}")
            return jsonify({"error": f"Failed to convert to DOCX: {str(e)}"}), 500
        
    except Exception as e:
        print(f"Admin save as DOC error: {e}")
        return jsonify({"error": f"Failed to save as DOC: {str(e)}"}), 500

@app.route('/api/admin/resume/<resume_id>/view', methods=['GET'])
def api_admin_view_resume(resume_id):
    """Admin route to view resume content by ID"""
    try:
        if 'admin' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        # Get resume from database by ID
        resume = db.get_resume_by_id(resume_id)
        if not resume:
            return jsonify({"error": "Resume not found"}), 404
        
        # Return extracted text content
        content = resume.get('extracted_text', 'No content available')
        candidate_name = resume.get('candidate_name', 'Unknown Candidate')
        filename = resume.get('filename', 'Unknown')
        
        return jsonify({
            "content": content,
            "candidate_name": candidate_name,
            "filename": filename
        }), 200
        
    except Exception as e:
        print(f"Admin view resume error: {e}")
        return jsonify({"error": "Failed to load resume content"}), 500

@app.route('/api/admin/user/<username>/remove', methods=['DELETE'])
def api_admin_remove_user(username):
    """Admin route to permanently remove a user"""
    try:
        if 'admin' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        # Prevent admin from removing themselves
        if username == session.get('admin'):
            return jsonify({"error": "Cannot remove your own admin account"}), 400
        
        # Remove user from database
        result = db.remove_user(username)
        
        if result:
            return jsonify({"message": f"User {username} removed successfully"}), 200
        else:
            return jsonify({"error": "User not found or failed to remove"}), 404
        
    except Exception as e:
        print(f"Admin remove user error: {e}")
        return jsonify({"error": "Failed to remove user"}), 500

@app.route('/api/admin/job-description', methods=['POST'])
def api_admin_add_job_description():
    """Admin route to add a new job description"""
    try:
        if 'admin' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        data = request.get_json()
        job_position = data.get('job_position', '').strip()
        job_title = data.get('job_title', '').strip()
        job_description = data.get('job_description', '').strip()
        
        if not job_position or not job_title or not job_description:
            return jsonify({"error": "All fields are required"}), 400
        
        # Save job description to database
        job_data = {
            "job_position": job_position,
            "job_title": job_title,
            "job_description": job_description,
            "created_by": session['admin'],
            "created_at": datetime.now(timezone.utc)
        }
        
        result = db.save_job_description(job_data)
        
        if result:
            return jsonify({"message": "Job description added successfully"}), 201
        else:
            return jsonify({"error": "Failed to save job description"}), 500
        
    except Exception as e:
        print(f"Admin add job description error: {e}")
        return jsonify({"error": "Failed to add job description"}), 500

@app.route('/api/admin/job-descriptions', methods=['GET'])
def api_admin_get_job_descriptions():
    """Admin route to get all job descriptions"""
    try:
        if 'admin' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        jobs = db.get_all_job_descriptions()
        
        # Convert ObjectId to string for JSON serialization
        for job in jobs:
            if '_id' in job:
                job['_id'] = str(job['_id'])
        
        return jsonify({"jobs": jobs}), 200
        
    except Exception as e:
        print(f"Admin get job descriptions error: {e}")
        return jsonify({"error": "Failed to fetch job descriptions"}), 500

@app.route('/api/admin/job-description/<job_id>', methods=['DELETE'])
def api_admin_delete_job_description(job_id):
    """Admin route to delete a job description"""
    try:
        if 'admin' not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        result = db.delete_job_description(job_id)
        
        if result:
            return jsonify({"message": "Job description deleted successfully"}), 200
        else:
            return jsonify({"error": "Job description not found"}), 404
        
    except Exception as e:
        print(f"Admin delete job description error: {e}")
        return jsonify({"error": "Failed to delete job description"}), 500

if __name__ == '__main__':
    import socket
    import sys
    
    # Check if port is available
    def is_port_available(port):
        try:
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                s.bind(('127.0.0.1', port))
                return True
        except OSError:
            return False
    
    # Find an available port
    port = 5000
    while not is_port_available(port) and port < 5010:
        port += 1
    
    if port >= 5010:
        print("❌ No available ports found between 5000-5009")
        sys.exit(1)
    
    print(f"🚀 Starting Flask app on port {port}")
    
    try:
        app.run(debug=True, host='127.0.0.1', port=port, threaded=True, use_reloader=False)
    except OSError as e:
        if "WinError 10038" in str(e):
            print("❌ Windows socket error detected. Trying alternative startup...")
            try:
                app.run(debug=False, host='127.0.0.1', port=port+1, threaded=True, use_reloader=False)
            except Exception as e2:
                print(f"❌ Failed to start server: {e2}")
                print("💡 Try running: python -m flask run --host=127.0.0.1 --port=5000")
        else:
            raise e